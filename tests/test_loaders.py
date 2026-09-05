import pytest

from assessment.documents import NeedsOCR, ParseError
from assessment.loaders import load_document


def test_markdown_preserves_code_table_and_locations(tmp_path):
    path = tmp_path / "guide.md"
    path.write_text("---\ntitle: hidden\n---\n# Transactions\n\nUse a transaction.\n\n```python\nif True:\n    commit()\n```\n\n| Key | Value |\n| --- | --- |\n| level | serializable |\n", encoding="utf-8")
    doc = load_document(path, source_uri="https://example.test/guide", version="v1")
    assert doc.title == "Transactions"
    assert "title: hidden" not in doc.text
    assert next(e for e in doc.elements if e.kind == "code").text == "if True:\n    commit()"
    assert "serializable" in next(e for e in doc.elements if e.kind == "table").text
    assert doc.elements[0].locator.line_start == 4


def test_html_removes_navigation_without_losing_table(tmp_path):
    path = tmp_path / "guide.html"
    path.write_text('<html><body><nav>Unrelated links</nav><main><h1 id="sql">SQL</h1><p>Use transactions.</p>'
                    '<pre>if ready:\n    commit()</pre><table><tr><th>Isolation</th></tr><tr><td>Serializable</td></tr></table>'
                    '<footer>Repeated footer</footer></main></body></html>', encoding="utf-8")
    doc = load_document(path)
    assert "Unrelated" not in doc.text and "footer" not in doc.text
    assert "    commit()" in doc.text
    assert "Isolation" in doc.text and "Serializable" in doc.text
    assert doc.elements[0].locator.anchor == "sql"


def test_docx_preserves_paragraph_table_order(tmp_path):
    from docx import Document

    document = Document()
    document.add_heading("Data pipelines", level=1)
    document.add_paragraph("First paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "Stage", "Action"
    table.cell(1, 0).text, table.cell(1, 1).text = "Load", "Validate"
    document.add_paragraph("Last paragraph.")
    path = tmp_path / "guide.docx"
    document.save(path)
    parsed = load_document(path)
    assert [e.kind for e in parsed.elements] == ["heading", "paragraph", "table", "paragraph"]
    assert "Load | Validate" in parsed.elements[2].text
    assert parsed.elements[2].locator.element == 2


def test_pdf_text_and_page_locator(tmp_path):
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=600, height=800)
    font = DictionaryObject({NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"),
                             NameObject("/BaseFont"): NameObject("/Helvetica")})
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})})
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 50 750 Td (Transactions preserve data consistency and isolate concurrent operations.) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(content)
    path = tmp_path / "guide.pdf"
    writer.write(path)
    document = load_document(path)
    assert "Transactions preserve" in document.text
    assert document.elements[0].locator.page == 1


def test_empty_scanned_pdf_is_not_silent_success(tmp_path):
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=600, height=800)
    path = tmp_path / "scan.pdf"
    writer.write(path)
    with pytest.raises(NeedsOCR):
        load_document(path)


def test_plain_text_encoding_and_invalid_input(tmp_path):
    path = tmp_path / "guide.txt"
    path.write_text("First paragraph.\n\nSecond paragraph.", encoding="utf-16")
    document = load_document(path)
    assert len(document.elements) == 2
    assert document.elements[1].locator.line_start == 3
    path.write_bytes(b"\x80\x81")
    with pytest.raises(ParseError):
        load_document(path)

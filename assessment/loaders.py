"""Format parsers preserve evidence structure; acquisition is a separate concern."""
from __future__ import annotations

import re
from pathlib import Path

from .documents import Document, Element, Locator, NeedsOCR, ParseError, digest

SUPPORTED = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx", ".rst"}


def decode(data: bytes) -> str:
    try:
        if data.startswith((b"\xff\xfe", b"\xfe\xff")):
            return data.decode("utf-16")
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ParseError("Text must be UTF-8 or BOM-marked UTF-16") from exc


def plain_elements(text: str, page: int | None = None) -> list[Element]:
    elements = []
    for match in re.finditer(r"\S[\s\S]*?(?=\n\s*\n|\Z)", text):
        value = match.group().rstrip()
        if not value:
            continue
        start = text.count("\n", 0, match.start()) + 1
        end = start + value.count("\n")
        elements.append(Element(kind="paragraph", text=value,
            locator=Locator(page=page, line_start=start, line_end=end, element=len(elements))))
    return elements


def markdown_elements(text: str) -> list[Element]:
    from markdown_it import MarkdownIt

    # Frontmatter and HTML comments are metadata, not retrieval evidence.
    text = re.sub(r"\A---\s*\n.*?\n---[^\n]*", lambda m: "\n" * m.group().count("\n"), text, flags=re.S)
    text = re.sub(r"<!--[\s\S]*?-->", lambda m: "\n" * m.group().count("\n"), text)
    tokens = MarkdownIt("commonmark").enable("table").parse(text)
    lines = text.splitlines()
    result, sections = [], []
    i = 0
    while i < len(tokens):
        token = tokens[i]
        value, kind = "", "paragraph"
        if token.type == "heading_open":
            value = tokens[i+1].content
            level = int(token.tag[1:])
            sections = sections[:level-1] + [value]
            kind = "heading"
        elif token.type in {"fence", "code_block"}:
            value, kind = token.content.rstrip(), "code"
        elif token.type == "paragraph_open":
            value = tokens[i+1].content
        elif token.type == "table_open":
            value = "\n".join(lines[slice(*token.map)])
            kind = "table"
            j = i+1
            while j < len(tokens) and tokens[j].type != "table_close":
                j += 1
            i = j
        if value.strip():
            span = token.map or [0, 0]
            result.append(Element(kind=kind, text=value, locator=Locator(
                line_start=span[0]+1, line_end=span[1], section=list(sections), element=len(result))))
        i += 1
    return result


def html_elements(data: bytes) -> list[Element]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "lxml")
    root = soup.select_one("main article, article.bd-article, main, [role=main], .document .body") or soup.body or soup
    for node in root.select("script, style, nav, footer, header, .headerlink, .toctree-wrapper, .related, .sphinxsidebar"):
        node.decompose()
    result, sections = [], []
    names = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "table", "li", "dt", "dd"}
    for node in root.find_all(list(names)):
        if any(parent.name in {"pre", "table", "li", "dd"} for parent in node.parents if parent is not root):
            continue
        kind = "paragraph"
        if node.name == "pre":
            kind, value = "code", node.get_text("", strip=False).rstrip()
        elif node.name == "table":
            kind = "table"
            value = "\n".join(" | ".join(cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"]))
                              for row in node.find_all("tr"))
        else:
            value = node.get_text(" ", strip=True)
            if re.fullmatch(r"h[1-6]", node.name):
                kind = "heading"
                sections = sections[:int(node.name[1])-1] + [value]
            elif node.name == "li":
                kind = "list"
        if not value.strip():
            continue
        parent_id = next((p.get("id") for p in node.parents if p.get("id")), None)
        result.append(Element(kind=kind, text=value, locator=Locator(
            anchor=node.get("id") or parent_id, element=len(result), section=list(sections))))
    if not result:
        result = plain_elements(root.get_text("\n", strip=True))
    return result


def pdf_elements(path: Path, ocr: bool) -> tuple[list[Element], list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted and not reader.decrypt(""):
        raise ParseError("Encrypted PDF requires a decrypted input")
    elements, warnings = [], []
    for number, page in enumerate(reader.pages, 1):
        text = (page.extract_text(extraction_mode="layout") or "") if "/Contents" in page else ""
        if sum(c.isalnum() for c in text) < 25:
            if not ocr:
                raise NeedsOCR(f"Page {number} has insufficient extractable text; enable OCR")
            import pypdfium2
            import pytesseract

            with pypdfium2.PdfDocument(path) as pdf:
                bitmap = pdf[number-1].render(scale=2.5)
                text = pytesseract.image_to_string(bitmap.to_pil(), lang="eng", timeout=45)
            if sum(c.isalnum() for c in text) < 25:
                raise ParseError(f"OCR produced insufficient text on page {number}")
            warnings.append(f"Page {number}: OCR text; review extraction accuracy")
        elements.extend(plain_elements(text, page=number))
    return elements, warnings


def docx_elements(path: Path) -> list[Element]:
    from docx import Document as WordDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = WordDocument(path)
    result, sections = [], []
    for position, item in enumerate(document.iter_inner_content()):
        kind = "paragraph"
        if isinstance(item, Paragraph):
            text = item.text.strip()
            style = item.style.name if item.style else ""
            if style.startswith("Heading"):
                kind = "heading"
                match = re.search(r"\d+", style)
                level = int(match[0]) if match else 1
                sections = sections[:level-1] + [text]
            elif "List" in style:
                kind = "list"
        elif isinstance(item, Table):
            kind = "table"
            text = "\n".join(" | ".join(cell.text for cell in row.cells) for row in item.rows)
        else:
            continue
        if text:
            result.append(Element(kind=kind, text=text,
                locator=Locator(element=position, section=list(sections))))
    return result


def load_document(path: Path, *, source_uri: str | None = None, version: str = "local",
                  license: str = "user-provided", ocr: bool = False) -> Document:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        raise ParseError(f"Unsupported document format: {suffix}")
    if path.stat().st_size > 256 * 1024 * 1024:
        raise ParseError("Document exceeds the 256 MiB parsing limit")
    raw = path.read_bytes()
    warnings = []
    try:
        if suffix in {".md", ".markdown"}:
            elements = markdown_elements(decode(raw))
        elif suffix in {".html", ".htm"}:
            elements = html_elements(raw)
        elif suffix == ".pdf":
            elements, warnings = pdf_elements(path, ocr)
        elif suffix == ".docx":
            elements = docx_elements(path)
        else:
            elements = plain_elements(decode(raw))
            if suffix == ".rst":
                warnings.append("RST ingested as source text; prefer a published HTML/TXT export")
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError(f"Failed to parse {suffix}: {type(exc).__name__}") from exc
    if not elements or not any(e.text.strip() for e in elements):
        raise ParseError("No usable document text was extracted")
    uri = source_uri or path.resolve().as_uri()
    title = next((e.text for e in elements if e.kind == "heading"), path.stem)
    return Document(source_id=digest(uri), source_uri=uri, version=version, title=title,
        format=suffix.lstrip("."), raw_hash=digest(raw), elements=elements, license=license, warnings=warnings)
